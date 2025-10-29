import io
from datetime import date
from typing import List, Dict

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Helpers ----------
def section_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def add_kv_line(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}:  ")
    run_label.bold = True
    run_value = p.add_run(value if value else "")
    run_value.bold = False

def add_bullets(doc: Document, items: List[str]):
    for it in items:
        if it.strip():
            doc.add_paragraph(it.strip(), style="List Number")

def make_docx(payload: Dict) -> bytes:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    tr = title.add_run("SHIP DAILY REPORT")
    tr.bold = True
    tr.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    # Header fields
    add_kv_line(doc, "SHIP", payload["ship"])
    add_kv_line(doc, "AVAILABILITY DATES", payload["avail_dates"])
    add_kv_line(doc, "LOCATION", payload["location"])
    add_kv_line(doc, "PPE", payload["ppe"])

    doc.add_paragraph("")  # spacer

    # Sections mirroring your example (with flexibility)
    for sec in payload["sections"]:
        section_title(doc, sec["title"])
        if sec["mode"] == "bullets":
            add_bullets(doc, sec["items"])
        else:
            # free text
            doc.add_paragraph(sec["text"])
        doc.add_paragraph("")  # spacer

    # Signature
    if payload["signature"]:
        section_title(doc, "V/r,")
        doc.add_paragraph(payload["signature"])

    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def make_txt(payload: Dict) -> str:
    lines = []
    lines.append("SHIP DAILY REPORT")
    lines.append("=" * 18)
    lines.append("")
    lines.append(f"SHIP: {payload['ship']}")
    lines.append(f"AVAILABILITY DATES: {payload['avail_dates']}")
    lines.append(f"LOCATION: {payload['location']}")
    lines.append(f"PPE: {payload['ppe']}")
    lines.append("")

    for sec in payload["sections"]:
        lines.append(f"{sec['title'].upper()}:")
        if sec["mode"] == "bullets":
            for i, it in enumerate([s for s in sec["items"] if s.strip()], start=1):
                lines.append(f"{i}. {it.strip()}")
        else:
            lines.append(sec["text"])
        lines.append("")

    if payload["signature"]:
        lines.append("V/r,")
        lines.append(payload["signature"])
        lines.append("")

    return "\n".join(lines)

def date_range_label(d1: date, d2: date) -> str:
    if d1 and d2:
        return f"{d1.strftime('%B %d, %Y')} – {d2.strftime('%B %d, %Y')}"
    return ""


def apply_theme(mode: str) -> None:
    themes = {
        "Light": {
            "color_scheme": "light",
            "background": "#f6f8fb",
            "sidebar": "#e4e7ec",
            "card": "#ffffff",
            "text": "#1f2933",
            "border": "#d0d7de",
            "accent": "#2563eb",
            "accent_border": "#1d4ed8",
            "button_text": "#ffffff",
            "code_bg": "#0f172a",
        },
        "Dark": {
            "color_scheme": "dark",
            "background": "#0f172a",
            "sidebar": "#111827",
            "card": "#1f2937",
            "text": "#f8fafc",
            "border": "#334155",
            "accent": "#38bdf8",
            "accent_border": "#0ea5e9",
            "button_text": "#03203c",
            "code_bg": "#020617",
        },
    }

    palette = themes.get(mode, themes["Light"])
    css = f"""
    <style>
    :root {{
        color-scheme: {palette['color_scheme']};
    }}

    body {{
        background-color: {palette['background']};
        color: {palette['text']};
    }}

    [data-testid="stAppViewContainer"] {{
        background-color: {palette['background']};
        color: {palette['text']};
    }}

    [data-testid="stHeader"] {{
        background: transparent;
    }}

    [data-testid="stSidebar"] > div:first-child {{
        background-color: {palette['sidebar']};
    }}

    .stMarkdown, .stCaption, .stText, .stExpander, .stCode {{
        color: {palette['text']};
    }}

    .stExpander {{
        background-color: {palette['card']};
        border: 1px solid {palette['border']};
        border-radius: 0.5rem;
    }}

    .stExpander:hover {{
        border-color: {palette['accent_border']};
    }}

    .stDownloadButton button, .stButton button {{
        background-color: {palette['accent']};
        color: {palette['button_text']};
        border: 1px solid {palette['accent_border']};
        border-radius: 0.5rem;
    }}

    .stDownloadButton button:hover, .stButton button:hover {{
        filter: brightness(0.95);
    }}

    textarea, input, .stTextInput > div > div > input {{
        background-color: {palette['card']};
        color: {palette['text']};
        border: 1px solid {palette['border']};
        border-radius: 0.5rem;
    }}

    textarea:focus, input:focus {{
        border-color: {palette['accent_border']};
        box-shadow: 0 0 0 1px {palette['accent_border']};
    }}

    .stTextArea textarea {{
        min-height: 160px;
    }}

    .stRadio label, .stCheckbox label {{
        color: {palette['text']};
    }}

    .stCode > div {{
        background-color: {palette['code_bg']};
        color: {palette['text']};
        border-radius: 0.5rem;
    }}

    </style>
    """
    st.markdown(css, unsafe_allow_html=True)


# ---------- UI ----------
st.set_page_config(page_title="Ship Daily Report Builder", page_icon="🛠️", layout="wide")
st.title("🛠️ Ship Daily Report Builder")

with st.sidebar:
    theme_mode = st.radio("Theme", ["Light", "Dark"], horizontal=True, index=0, key="theme_mode")
    st.markdown("---")
    st.markdown("**Report Header**")
    ship = st.text_input("Ship Name", placeholder="e.g., USNS Comfort (T-AH-20)")
    colA, colB = st.columns(2)
    with colA:
        start_date = st.date_input("Availability Start", value=None)
    with colB:
        end_date = st.date_input("Availability End", value=None)
    location = st.text_input("Location", placeholder="e.g., Alabama Shipyard, Mobile, Alabama")
    ppe = st.text_input("PPE", placeholder="e.g., Steven Destree")

    st.markdown("---")
    st.markdown("**Signature**")
    signature = st.text_input("Sign-off (optional)", placeholder="e.g., Michael McCary, Robert Tippitt")

apply_theme(theme_mode)

st.markdown("### Sections")
st.caption("Add content below. Use the toggles to choose bullet lists or free text for each section.")

DEFAULT_SECTIONS = [
    ("WORK SUMMARY", True),
    ("PAINT PROGRESS", True),
    ("DAILY ACTIVITY", True),
    ("ITEMS OF INTEREST", True),
]

# Build dynamic sections
sections_data = []
for idx, (title, bullets_default) in enumerate(DEFAULT_SECTIONS, start=1):
    with st.expander(f"{idx}. {title}", expanded=(idx <= 2)):
        mode = st.radio(
            f"Input mode for {title}",
            options=["Bulleted Items", "Free Text"],
            key=f"mode_{title}",
            horizontal=True,
            index=0 if bullets_default else 1
        )
        if mode == "Bulleted Items":
            st.caption("Enter one item per line. Numbers are added automatically.")
            items_text = st.text_area(
                f"{title} items",
                key=f"items_{title}",
                placeholder="1) WI#0021 ... 95%\n2) WI#0022 ... 5%\n"
            )
            items = [line.strip(" \t").lstrip("0123456789). ").strip() for line in items_text.splitlines() if line.strip()]
            sections_data.append({"title": title, "mode": "bullets", "items": items, "text": ""})
        else:
            text_val = st.text_area(f"{title} text", key=f"text_{title}", height=150, placeholder=f"Describe {title.lower()} here…")
            sections_data.append({"title": title, "mode": "text", "items": [], "text": text_val})

# Build payload
avail_dates = date_range_label(start_date, end_date) if (start_date and end_date and start_date <= end_date) else ""
payload = {
    "ship": ship.strip(),
    "avail_dates": avail_dates,
    "location": location.strip(),
    "ppe": ppe.strip(),
    "signature": signature.strip(),
    "sections": sections_data,
}

# Preview
st.markdown("---")
st.subheader("Live Preview")
preview_txt = make_txt(payload)
st.code(preview_txt, language="markdown")

# Downloads
st.markdown("### Export")
col1, col2 = st.columns(2)

with col1:
    if st.button("Generate .docx"):
        doc_bytes = make_docx(payload)
        st.download_button(
            label="Download Daily Report (.docx)",
            data=doc_bytes,
            file_name=f"{(ship or 'Ship').replace(' ', '_')}_Daily_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

with col2:
    st.download_button(
        label="Download Daily Report (.txt)",
        data=preview_txt.encode("utf-8"),
        file_name=f"{(ship or 'Ship').replace(' ', '_')}_Daily_Report.txt",
        mime="text/plain",
        use_container_width=True
    )

st.markdown("---")
st.caption("Tip: Keep a shared folder for the exported reports so the team can find them easily.")
